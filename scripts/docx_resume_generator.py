from docx import Document
import os


def generate_docx_resume(
        job_title,
        company,
        summary,
        skills
):

    os.makedirs(
        "generated_docx_resumes",
        exist_ok=True
    )

    doc = Document()

    doc.add_heading(
        "Harish MR",
        level=1
    )

    doc.add_paragraph(
        "SOC Analyst | Network Engineer | ISO 27001 Lead Auditor"
    )

    doc.add_heading(
        "Professional Summary",
        level=2
    )

    doc.add_paragraph(
        summary
    )

    doc.add_heading(
        "Key Skills",
        level=2
    )

    if isinstance(skills, list):
        skills_text = ", ".join(skills)
    else:
        skills_text = skills

    doc.add_paragraph(
        skills_text
    )

    doc.add_heading(
        "Experience",
        level=2
    )

    doc.add_paragraph(
        """
Network Engineer
Arridae Infosec Pvt Ltd

• Network Monitoring using NMIS
• SIEM Monitoring
• Linux Administration
• Firewall Management
• DNS and DHCP Troubleshooting
• Routing and Switching
• VLAN Configuration
• Security Operations Support
"""
    )

    doc.add_heading(
        "Target Position",
        level=2
    )

    doc.add_paragraph(
        f"{job_title} at {company}"
    )

    filename = (
        f"generated_docx_resumes/"
        f"{company}_{job_title}.docx"
    )

    filename = (
        filename
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "_")
    )

    doc.save(
        filename
    )

    print(
        f"DOCX Resume Generated: {filename}"
    )