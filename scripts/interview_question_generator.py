from docx import Document
import os


def generate_interview_questions(
        job_title,
        company
):

    os.makedirs(
        "interview_prep",
        exist_ok=True
    )

    questions = [

        f"Tell me about yourself.",

        f"Why do you want to join {company}?",

        "Explain TCP/IP.",

        "What is DNS?",

        "What is DHCP?",

        "Difference between TCP and UDP?",

        "What is VLAN?",

        "What is Inter-VLAN Routing?",

        "What is STP?",

        "Explain OSPF.",

        "What is a Firewall?",

        "What is SIEM?",

        "Explain Incident Response.",

        "What is a Security Incident?",

        "What is Linux?",

        "Explain Active Directory.",

        "What is Wireshark?",

        "What is Nmap?",

        "How do you troubleshoot network issues?",

        "Explain your current role and responsibilities."
    ]

    doc = Document()

    doc.add_heading(
        f"{job_title} Interview Questions",
        level=1
    )

    doc.add_paragraph(
        f"Company: {company}"
    )

    for i, q in enumerate(
        questions,
        start=1
    ):

        doc.add_paragraph(
            f"{i}. {q}"
        )

    filename = (
        f"interview_prep/"
        f"{company}_{job_title}_Questions.docx"
    )

    filename = (
        filename
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "_")
    )

    doc.save(
        filename
    )

    print(
        f"Interview Questions Generated: {filename}"
    )