from docx import Document
import os


def generate_docx_cover_letter(
        company,
        job_title,
        skills
):

    os.makedirs(
        "generated_docx_cover_letters",
        exist_ok=True
    )

    if isinstance(skills, list):
        skills = ", ".join(skills)

    doc = Document()

    doc.add_heading(
        "Cover Letter",
        level=1
    )

    letter = f"""
Dear Hiring Manager,

I am writing to express my interest in the {job_title} position at {company}.

I have hands-on experience in Networking, Cyber Security,
SIEM Monitoring, Linux Administration, Windows Administration,
Firewall Management, DNS, DHCP, Routing and Switching.

My technical skills include:
{skills}

I currently work as a Network Engineer and support enterprise
network infrastructure, monitoring, troubleshooting,
security operations and incident management activities.

In addition, I have developed automation and AI-based projects
using Python, SQLite, APIs and Streamlit which strengthened
my analytical and problem-solving abilities.

I am excited about the opportunity to contribute my skills
and continue growing within your organization.

Thank you for your time and consideration.

Sincerely,

Harish MR
"""

    doc.add_paragraph(
        letter
    )

    filename = (
        f"generated_docx_cover_letters/"
        f"{company}_{job_title}.docx"
    )

    filename = (
        filename
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "_")
    )

    doc.save(
        filename
    )

    print(
        f"DOCX Cover Letter Generated: {filename}"
    )