from docx import Document
import os


def generate_interview_answers(
        job_title,
        company
):

    os.makedirs(
        "interview_answers",
        exist_ok=True
    )

    qa = {

        "What is SIEM?":
        """
SIEM (Security Information and Event Management)
is a security solution that collects, correlates,
and analyzes logs from various systems and devices
to identify security incidents and suspicious activities.
        """,

        "What is DNS?":
        """
DNS (Domain Name System) translates domain names
into IP addresses so that computers can communicate
over a network.
        """,

        "What is DHCP?":
        """
DHCP (Dynamic Host Configuration Protocol)
automatically assigns IP addresses and network
configuration parameters to devices.
        """,

        "Difference between TCP and UDP?":
        """
TCP is connection-oriented and reliable,
while UDP is connectionless and faster but
does not guarantee delivery.
        """,

        "What is VLAN?":
        """
A VLAN (Virtual Local Area Network)
logically separates devices within the same
physical network to improve security and management.
        """,

        "What is a Firewall?":
        """
A Firewall monitors and controls incoming
and outgoing network traffic based on security rules.
        """,

        "What is Incident Response?":
        """
Incident Response is the process of identifying,
investigating, containing, eradicating and recovering
from security incidents.
        """,

        "Explain Linux.":
        """
Linux is an open-source operating system widely used
for servers, networking, cybersecurity and cloud environments.
        """,

        "What is Wireshark?":
        """
Wireshark is a packet analysis tool used to capture
and inspect network traffic for troubleshooting
and security investigations.
        """,

        "What is Nmap?":
        """
Nmap is a network scanning tool used to discover hosts,
services, open ports and security vulnerabilities.
        """
    }

    doc = Document()

    doc.add_heading(
        f"{job_title} Interview Answers",
        level=1
    )

    doc.add_paragraph(
        f"Company: {company}"
    )

    for question, answer in qa.items():

        doc.add_heading(
            question,
            level=2
        )

        doc.add_paragraph(
            answer.strip()
        )

    filename = (
        f"interview_answers/"
        f"{company}_{job_title}_Answers.docx"
    )

    filename = (
        filename
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "_")
    )

    doc.save(
        filename
    )

    print(
        f"Interview Answers Generated: {filename}"
    )