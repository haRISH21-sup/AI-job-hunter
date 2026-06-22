from docx import Document
import os


def analyze_skill_gap(
        company,
        job_title,
        resume_skills,
        job_description
):

    os.makedirs(
        "skill_gap_reports",
        exist_ok=True
    )

    keywords = [

        "SIEM",
        "SOC",
        "Linux",
        "Windows",
        "Python",
        "Networking",
        "Firewall",
        "TCP/IP",
        "DNS",
        "DHCP",
        "Wireshark",
        "Nmap",
        "Incident Response",
        "Splunk",
        "QRadar",
        "Azure",
        "AWS",
        "Active Directory",
        "Routing",
        "Switching",
        "Cybersecurity"
    ]

    matched = []
    missing = []

    job_lower = job_description.lower()

    for skill in keywords:

        if skill.lower() in job_lower:

            if skill in resume_skills:
                matched.append(skill)
            else:
                missing.append(skill)

    doc = Document()

    doc.add_heading(
        "Skill Gap Analysis",
        level=1
    )

    doc.add_paragraph(
        f"Company: {company}"
    )

    doc.add_paragraph(
        f"Role: {job_title}"
    )

    doc.add_heading(
        "Matched Skills",
        level=2
    )

    if matched:

        for item in matched:
            doc.add_paragraph(
                f"✓ {item}"
            )

    else:

        doc.add_paragraph(
            "No direct matches found."
        )

    doc.add_heading(
        "Missing Skills",
        level=2
    )

    if missing:

        for item in missing:
            doc.add_paragraph(
                f"✗ {item}"
            )

    else:

        doc.add_paragraph(
            "No major skill gaps detected."
        )

    doc.add_heading(
        "Recommended Learning Areas",
        level=2
    )

    if missing:

        for item in missing:
            doc.add_paragraph(
                f"Learn: {item}"
            )

    else:

        doc.add_paragraph(
            "Resume aligns well with this role."
        )

    filename = (
        f"skill_gap_reports/"
        f"{company}_{job_title}_SkillGap.docx"
    )

    filename = (
        filename
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "_")
    )

    doc.save(
        filename
    )

    print(
        f"Skill Gap Report Generated: {filename}"
    )