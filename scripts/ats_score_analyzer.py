from docx import Document
import os


def generate_ats_report(
        company,
        job_title,
        resume_skills,
        job_description,
        match_score
):

    os.makedirs(
        "ats_reports",
        exist_ok=True
    )

    keywords = [

        "SIEM",
        "SOC",
        "Linux",
        "Windows",
        "Python",
        "Networking",
        "Firewall",
        "TCP/IP",
        "DNS",
        "DHCP",
        "Wireshark",
        "Nmap",
        "Incident Response",
        "Splunk",
        "QRadar",
        "Azure",
        "AWS",
        "Active Directory",
        "Routing",
        "Switching",
        "Cybersecurity"
    ]

    matched = []
    missing = []

    job_lower = job_description.lower()

    for skill in keywords:

        if skill.lower() in job_lower:

            if skill in resume_skills:
                matched.append(skill)
            else:
                missing.append(skill)

    doc = Document()

    doc.add_heading(
        "ATS Resume Analysis Report",
        level=1
    )

    doc.add_paragraph(
        f"Company: {company}"
    )

    doc.add_paragraph(
        f"Role: {job_title}"
    )

    doc.add_paragraph(
        f"ATS Match Score: {round(match_score,2)}%"
    )

    doc.add_heading(
        "Matched Keywords",
        level=2
    )

    if matched:

        for item in matched:
            doc.add_paragraph(
                f"✓ {item}"
            )

    else:

        doc.add_paragraph(
            "No matched keywords found."
        )

    doc.add_heading(
        "Missing Keywords",
        level=2
    )

    if missing:

        for item in missing:
            doc.add_paragraph(
                f"✗ {item}"
            )

    else:

        doc.add_paragraph(
            "No major missing keywords."
        )

    doc.add_heading(
        "Recommendations",
        level=2
    )

    if missing:

        for item in missing:
            doc.add_paragraph(
                f"Add experience, projects or skills related to {item}"
            )

    else:

        doc.add_paragraph(
            "Resume is well aligned with this role."
        )

    filename = (
        f"ats_reports/"
        f"{company}_{job_title}_ATS_Report.docx"
    )

    filename = (
        filename
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "_")
    )

    doc.save(
        filename
    )

    print(
        f"ATS Report Generated: {filename}"
    )